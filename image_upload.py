from flask import Flask, request, render_template_string, send_file
from PIL import Image
import pytesseract
from docx import Document
import os
from werkzeug.utils import secure_filename
import uuid
import time
import io

# Configure Tesseract-OCR path (Windows-specific)
pytesseract.pytesseract.tesseract_cmd = r'C:/Program Files/Tesseract-OCR/tesseract.exe'

app = Flask(__name__)
UPLOAD_FOLDER = 'uploads'
DOCS_FOLDER = 'docs'
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['DOCS_FOLDER'] = DOCS_FOLDER
app.config['MAX_CONTENT_LENGTH'] = 5 * 1024 * 1024  # 5MB max file size
ALLOWED_EXTENSIONS = {'png', 'jpg', 'jpeg', 'gif'}

# Create folders if they don't exist
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(DOCS_FOLDER, exist_ok=True)

def allowed_file(filename):
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

@app.route('/', methods=['GET', 'POST'])
def upload_image():
    error = None
    text = None
    image_url = None
    doc_filename = None

    if request.method == 'POST':
        if 'image' not in request.files:
            error = 'No file uploaded.'
        else:
            file = request.files['image']
            if file.filename == '':
                error = 'No file selected.'
            elif not allowed_file(file.filename):
                error = 'Invalid file type. Only PNG, JPGdemand, JPEG, and GIF are allowed.'
            else:
                # Generate unique filename for image
                filename = secure_filename(file.filename)
                unique_filename = f"{uuid.uuid4().hex}_{int(time.time())}.{filename.rsplit('.', 1)[1].lower()}"
                filepath = os.path.join(app.config['UPLOAD_FOLDER'], unique_filename)
                
                try:
                    # Save and process image
                    file.save(filepath)
                    image = Image.open(filepath)
                    text = pytesseract.image_to_string(image)
                    image_url = f"/{filepath}"

                    # Generate Word document
                    doc = Document()
                    doc.add_heading('OCR Extracted Text', 0)
                    doc.add_paragraph(text)
                    doc_filename = f"ocr_output_{unique_filename}.docx"
                    doc_filepath = os.path.join(app.config['DOCS_FOLDER'], doc_filename)
                    doc.save(doc_filepath)

                except Exception as e:
                    error = f"Error processing image: {str(e)}"
                finally:
                    # Clean up image file
                    if os.path.exists(filepath):
                        os.remove(filepath)

    # Unified template for upload and result
    template = '''
    <!DOCTYPE html>
    <html lang="en">
    <head>
        <meta charset="UTF-8">
        <meta name="viewport" content="width=device-width, initial-scale=1.0">
        <title>Upload an Image for OCR</title>
        <style>
            * {
                margin: 0;
                padding: 0;
                box-sizing: border-box;
                font-family: 'Poppins', sans-serif;
            }

            body {
                display: flex;
                justify-content: center;
                align-items: center;
                min-height: 100vh;
                background: linear-gradient(135deg, #6e8efb, #a777e3);
                color: #fff;
                padding: 20px;
            }

            .container {
                background: rgba(255, 255, 255, 0.1);
                backdrop-filter: blur(10px);
                border-radius: 15px;
                padding: 40px;
                max-width: 600px;
                width: 100%;
                text-align: center;
                box-shadow: 0 8px 32px rgba(0, 0, 0, 0.2);
                animation: fadeIn 1s ease-in-out;
            }

            @keyframes fadeIn {
                from { opacity: 0; transform: translateY(-20px); }
                to { opacity: 1; transform: translateY(0); }
            }

            h1 {
                font-size: 2.2em;
                margin-bottom: 20px;
                color: #fff;
                text-shadow: 0 2px 4px rgba(0, 0, 0, 0.2);
            }

            .upload-area {
                border: 2px dashed #fff;
                border-radius: 10px;
                padding: 20px;
                margin-bottom: 20px;
                cursor: pointer;
                transition: all 0.3s ease;
            }

            .upload-area:hover, .upload-area.dragover {
                background: rgba(255, 255, 255, 0.2);
                border-color: #a777e3;
            }

            .upload-area p {
                font-size: 1.1em;
                color: #ddd;
            }

            input[type="file"] {
                display: none;
            }

            .custom-file-label {
                display: inline-block;
                padding: 10px 20px;
                background: #a777e3;
                color: #fff;
                border-radius: 25px;
                cursor: pointer;
                transition: background 0.3s ease;
                margin-top: 10px;
            }

            .custom-file-label:hover {
                background: #8c5ed2;
            }

            input[type="submit"] {
                background: #6e8efb;
                border: none;
                padding: 12px 30px;
                color: #fff;
                font-size: 1.1em;
                border-radius: 25px;
                cursor: pointer;
                transition: all 0.3s ease;
                margin-top: 10px;
            }

            input[type="submit"]:hover {
                background: #5a75e3;
                transform: translateY(-2px);
                box-shadow: 0 4px 15px rgba(0, 0, 0, 0.2);
            }

            input[type="submit"]:active {
                transform: translateY(0);
            }

            .file-name {
                margin-top: 10px;
                font-size: 0.9em;
                color: #ddd;
            }

            .error {
                color: #ff6b6b;
                margin-bottom: 20px;
                font-size: 1em;
            }

            .result {
                margin-top: 20px;
                text-align: left;
            }

            .result img {
                max-width: 100%;
                border-radius: 10px;
                margin-bottom: 20px;
            }

            .result pre {
                background: rgba(0, 0, 0, 0.2);
                padding: 15px;
                border-radius: 10px;
                white-space: pre-wrap;
                max-height: 300px;
                overflow-y: auto;
            }

            .back-btn, .download-btn {
                display: inline-block;
                padding: 10px 20px;
                color: #fff;
                text-decoration: none;
                border-radius: 25px;
                margin-top: 20px;
                transition: background 0.3s ease;
            }

            .back-btn {
                background: #ff6b6b;
                margin-right: 10px;
            }

            .back-btn:hover {
                background: #e55a5a;
            }

            .download-btn {
                background: #28c76f;
            }

            .download-btn:hover {
                background: #24b363;
            }

            .loading {
                display: none;
                font-size: 1em;
                color: #ddd;
                margin-top: 10px;
            }

            @media (max-width: 600px) {
                .container {
                    padding: 20px;
                }

                h1 {
                    font-size: 1.8em;
                }

                .upload-area p {
                    font-size: 1em;
                }
            }
        </style>
    </head>
    <body>
        <div class="container">
            {% if text or error %}
                <h1>OCR Result</h1>
                {% if error %}
                    <p class="error">{{ error }}</p>
                {% endif %}
                {% if text %}
                    <div class="result">
                        {% if image_url %}
                            <img src="{{ image_url }}" alt="Uploaded Image">
                        {% endif %}
                        <h2>Recognized Text:</h2>
                        <pre>{{ text }}</pre>
                    </div>
                    <a href="/" class="back-btn">Upload Another Image</a>
                    {% if doc_filename %}
                        <a href="/download/{{ doc_filename }}" class="download-btn">Download as Word</a>
                    {% endif %}
                {% endif %}
            {% else %}
                <h1>Upload an Image for OCR</h1>
                {% if error %}
                    <p class="error">{{ error }}</p>
                {% endif %}
                <form method="post" enctype="multipart/form-data" id="uploadForm">
                    <div class="upload-area" id="uploadArea">
                        <p>Drag & Drop your image here or click to select</p>
                        <label for="image" class="custom-file-label">Choose File</label>
                        <input type="file" name="image" id="image" accept="image/*">
                        <div class="file-name" id="fileName"></div>
                    </div>
                    <input type="submit" value="Upload" id="submitBtn">
                    <div class="loading" id="loading">Processing...</div>
                </form>
            {% endif %}
        </div>

        <script>
            const uploadArea = document.getElementById('uploadArea');
            const fileInput = document.getElementById('image');
            const fileNameDisplay = document.getElementById('fileName');
            const form = document.getElementById('uploadForm');
            const submitBtn = document.getElementById('submitBtn');
            const loading = document.getElementById('loading');

            uploadArea.addEventListener('click', () => fileInput.click());

            uploadArea.addEventListener('dragover', (e) => {
                e.preventDefault();
                uploadArea.classList.add('dragover');
            });

            uploadArea.addEventListener('dragleave', () => {
                uploadArea.classList.remove('dragover');
            });

            uploadArea.addEventListener('drop', (e) => {
                e.preventDefault();
                uploadArea.classList.remove('dragover');
                const files = e.dataTransfer.files;
                if (files.length) {
                    fileInput.files = files;
                    fileNameDisplay.textContent = files[0].name;
                }
            });

            fileInput.addEventListener('change', () => {
                if (fileInput.files.length) {
                    fileNameDisplay.textContent = fileInput.files[0].name;
                } else {
                    fileNameDisplay.textContent = '';
                }
            });

            form.addEventListener('submit', () => {
                submitBtn.disabled = true;
                loading.style.display = 'block';
            });
        </script>
    </body>
    </html>
    '''
    return render_template_string(template, error=error, text=text, image_url=image_url, doc_filename=doc_filename)

@app.route('/download/<filename>')
def download_file(filename):
    filepath = os.path.join(app.config['DOCS_FOLDER'], filename)
    try:
        response = send_file(filepath, as_attachment=True)
        # Clean up the .docx file after sending
        os.remove(filepath)
        return response
    except FileNotFoundError:
        return "File not found", 404

if __name__ == '__main__':
    app.run(debug=True)